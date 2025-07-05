import PyPDF2
from docx import Document
import re
from typing import Dict, List
import os

class ResumeParser:
    def __init__(self):
        # Comprehensive skills database
        self.skills_keywords = [
            # Programming Languages
            'python', 'javascript', 'java', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift',
            'kotlin', 'scala', 'r', 'matlab', 'sql', 'typescript', 'dart', 'objective-c',
            
            # Web Technologies
            'html', 'css', 'react', 'angular', 'vue', 'svelte', 'jquery', 'bootstrap',
            'tailwind', 'sass', 'less', 'webpack', 'vite', 'next.js', 'nuxt.js',
            
            # Backend & APIs
            'node.js', 'express', 'fastapi', 'django', 'flask', 'spring', 'spring boot',
            'asp.net', 'laravel', 'rails', 'gin', 'fiber', 'rest api', 'graphql',
            
            # Databases
            'mongodb', 'mysql', 'postgresql', 'sqlite', 'redis', 'elasticsearch',
            'cassandra', 'dynamodb', 'oracle', 'sql server', 'firebase',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'gitlab ci',
            'github actions', 'terraform', 'ansible', 'chef', 'puppet',
            
            # Data Science & AI
            'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'scikit-learn',
            'pandas', 'numpy', 'matplotlib', 'seaborn', 'jupyter', 'tableau', 'power bi',
            
            # Mobile Development
            'android', 'ios', 'react native', 'flutter', 'xamarin', 'ionic',
            
            # Tools & Others
            'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'slack',
            'linux', 'ubuntu', 'windows', 'macos', 'bash', 'powershell',
            
            # Methodologies
            'agile', 'scrum', 'kanban', 'devops', 'ci/cd', 'tdd', 'bdd',
            'microservices', 'api design', 'system design'
        ]
        
        # Alternative skill names mapping
        self.skill_aliases = {
            'js': 'javascript',
            'ts': 'typescript',
            'nodejs': 'node.js',
            'reactjs': 'react',
            'angularjs': 'angular',
            'vuejs': 'vue',
            'py': 'python',
            'ml': 'machine learning',
            'ai': 'artificial intelligence',
            'db': 'database',
            'api': 'rest api',
            'ui': 'user interface',
            'ux': 'user experience'
        }
    
    def extract_text(self, file_path: str) -> Dict:
        """Extract text from PDF or DOCX file and parse resume content"""
        try:
            print(f"🔍 Processing file: {file_path}")
            
            if file_path.lower().endswith('.pdf'):
                text = self._extract_from_pdf(file_path)
            elif file_path.lower().endswith(('.docx', '.doc')):
                text = self._extract_from_docx(file_path)
            else:
                raise ValueError("Unsupported file format")
            
            print(f"📝 Extracted text length: {len(text)} characters")
            
            parsed_data = self._parse_resume_content(text)
            print(f"🎯 Found {len(parsed_data['skills'])} skills")
            
            return parsed_data
            
        except Exception as e:
            print(f"❌ Error extracting text: {str(e)}")
            raise Exception(f"Error extracting text: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    print(f"📄 Page {page_num + 1}: {len(page_text)} characters")
                    
        except Exception as e:
            print(f"❌ PDF extraction error: {str(e)}")
            raise Exception(f"Error reading PDF: {str(e)}")
        
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(file_path)
            for para_num, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                        
        except Exception as e:
            print(f"❌ DOCX extraction error: {str(e)}")
            raise Exception(f"Error reading DOCX: {str(e)}")
        
        return text
    
    def _parse_resume_content(self, text: str) -> Dict:
        """Parse resume content and extract key information"""
        text_lower = text.lower()
        
        # Extract skills with better matching
        found_skills = self._extract_skills(text_lower)
        
        # Extract contact information
        email = self._extract_email(text)
        phone = self._extract_phone(text)
        
        # Extract experience
        experience_years = self._extract_experience(text_lower)
        
        # Extract education
        education = self._extract_education(text_lower)
        
        # Create better summary
        summary = self._create_summary(text)
        
        return {
            "full_text": text,
            "summary": summary,
            "skills": found_skills,
            "experience_years": experience_years,
            "email": email,
            "phone": phone,
            "education": education,
            "word_count": len(text.split())
        }
    
    def _extract_skills(self, text_lower: str) -> List[str]:
        """Extract skills with improved matching"""
        found_skills = set()
        
        # Direct skill matching
        for skill in self.skills_keywords:
            if skill.lower() in text_lower:
                found_skills.add(skill.title())
        
        # Alias matching
        for alias, skill in self.skill_aliases.items():
            if alias.lower() in text_lower:
                found_skills.add(skill.title())
        
        # Pattern-based matching for common formats
        skill_patterns = [
            r'skills?[:\s]*([^.]+)',
            r'technologies?[:\s]*([^.]+)',
            r'programming languages?[:\s]*([^.]+)',
            r'tools?[:\s]*([^.]+)',
            r'frameworks?[:\s]*([^.]+)'
        ]
        
        for pattern in skill_patterns:
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                # Split by common delimiters
                potential_skills = re.split(r'[,;|•\n\r]', match)
                for skill in potential_skills:
                    skill = skill.strip()
                    if skill and len(skill) > 1:
                        # Check if it's a known skill
                        for known_skill in self.skills_keywords:
                            if known_skill.lower() in skill.lower():
                                found_skills.add(known_skill.title())
        
        return list(found_skills)
    
    def _extract_email(self, text: str) -> str:
        """Extract email address"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else None
    
    def _extract_phone(self, text: str) -> str:
        """Extract phone number"""
        phone_patterns = [
            r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',
            r'\+?[1-9][0-9]{7,14}',
            r'[\+]?[1-9]?[0-9]{7,14}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0]
        return None
    
    def _extract_experience(self, text_lower: str) -> int:
        """Extract years of experience"""
        experience_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'(\d+)\+?\s*yrs?\s*(?:of\s*)?experience',
            r'experience[:\s]*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s*in\s*(?:software|programming|development)',
            r'over\s*(\d+)\s*years?',
            r'more\s*than\s*(\d+)\s*years?'
        ]
        
        years = []
        for pattern in experience_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                try:
                    years.append(int(match))
                except ValueError:
                    continue
        
        return max(years) if years else 0
    
    def _extract_education(self, text_lower: str) -> List[str]:
        """Extract education information"""
        education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'degree', 'diploma',
            'b.tech', 'm.tech', 'b.sc', 'm.sc', 'mca', 'bca', 'mba',
            'university', 'college', 'institute', 'graduation', 'undergraduate',
            'postgraduate', 'certification', 'certificate'
        ]
        
        found_education = set()
        for edu in education_keywords:
            if edu in text_lower:
                found_education.add(edu.title())
        
        return list(found_education)
    
    def _create_summary(self, text: str) -> str:
        """Create a better summary from resume text"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Look for summary/objective sections
        summary_indicators = ['summary', 'objective', 'profile', 'about']
        summary_text = ""
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(indicator in line_lower for indicator in summary_indicators):
                # Take next few lines as summary
                for j in range(i + 1, min(i + 6, len(lines))):
                    if lines[j] and not lines[j].isupper():
                        summary_text += lines[j] + " "
                break
        
        # If no summary section found, use first few meaningful lines
        if not summary_text:
            meaningful_lines = [line for line in lines[:10] 
                              if len(line) > 20 and not line.isupper()]
            summary_text = " ".join(meaningful_lines[:3])
        
        return summary_text.strip()[:500]  # Limit to 500 characters